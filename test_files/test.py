# Create the Word document based on the content from the canvas
from docx import Document

# Initialize a new Document
doc = Document()

# Add title and content to the Word document
doc.add_heading("Deconoix Gloire Gabell", level=1)

doc.add_heading("PROFILE", level=2)
doc.add_paragraph(
    "Motivated software engineering student at MSU Denver studying Computer Science with hands-on experience in web and Python development. "
    "Inducted member and acting president of the National Society of Leadership and Success. "
    "Adept at designing, troubleshooting, and optimizing technical solutions. Skilled in managing complex projects, solving technical problems, "
    "and collaborating within multicultural settings. Passionate about learning new technologies and delivering innovative solutions."
)

doc.add_heading("CONTACT", level=2)
doc.add_paragraph("303-908-7207")
doc.add_paragraph("Deconoixggabell@gmail.com")
doc.add_paragraph("LinkedIn Profile: https://www.linkedin.com/in/deconoix-gloire-gabell-04b04b1a3")
doc.add_paragraph("GitHub Profile: https://github.com/deconoixggabell")

doc.add_heading("SKILLS", level=2)
doc.add_paragraph(
    "- Programming: Python, HTML, CSS, C/C++, SQL\n"
    "- Frameworks & Tools: Active Directory, VMware, Jira, WinSCP, Putty\n"
    "- Operating Systems: Linux (Arch, Ubuntu), Windows OS (7, 8, 10, 11)\n"
    "- Networking: Cisco equipment, Routers, IP Cameras, NTP Server\n"
    "- Database: SQL, Microsoft Server (2012, 2014, 2016, 2019)\n"
    "- Other Tools: Microsoft Visio, 3D Printers, Laser Cutters, Arduino\n"
    "- Languages: Bilingual (English and French)\n"
    "- Soft Skills: Handling Confidential Information, Problem-solving, Multicultural teamwork"
)

doc.add_heading("EDUCATION", level=2)
doc.add_paragraph(
    "- Bachelor of Science in Computer Science\n"
    "  Metropolitan State University of Denver, Denver, CO\n"
    "  Expected Graduation: [Insert Expected Graduation Year]\n\n"
    "- Associate of Science in Computer Science\n"
    "  Red Rocks Community College, Lakewood, CO\n"
    "  2016 – 2020\n\n"
    "- Certification in Software Development\n"
    "  Coding Dojo"
)

doc.add_heading("EXPERIENCE", level=2)
doc.add_paragraph(
    "- SKIDATA INC – Event System Engineer\n"
    "  February 2020 – February 2022\n"
    "  - Managed technical issues for customers nationwide, providing onsite and remote support.\n"
    "  - Coordinated and led installation teams for new equipment setup.\n"
    "  - Resolved Linux-based system issues and maintained SQL databases.\n"
    "  - Delivered exceptional customer service, ensuring smooth system operations.\n"
    "  - Key Achievements: Enhanced troubleshooting processes, improving system uptime by 20%.\n\n"
    "- Red Rocks Community College IT Department\n"
    "  Helpdesk Supervisor (December 2017 – May 2018)\n"
    "  - Resolved technical issues for students and staff, documenting and managing tickets via Jira.\n"
    "  - Led weekly team meetings and ensured efficient scheduling.\n"
    "  - Specialized in credential management, password resets, and troubleshooting Active Directory.\n\n"
    "  System Engineer Assistant (March 2017 – May 2020)\n"
    "  - Transitioned physical machines to a remote VMware environment.\n"
    "  - Collaborated on large-scale projects, including system migrations and network setups.\n\n"
    "- Oracle Security Analyst\n"
    "  August 2018 – February 2020\n"
    "  - Monitored facilities nationwide to ensure system security.\n"
    "  - Programmed and managed door access systems.\n\n"
    "- Robotics Engineer – Idea Lab\n"
    "  January 2019 – August 2020\n"
    "  - Assisted engineering students with projects and homework.\n"
    "  - Designed and built a seven-headed dragon for a theater production.\n"
    "  - Programmed Arduino boards, built circuit boards, and managed 3D printing and laser cutting.\n\n"
    "- Gabell Cleaning Manager\n"
    "  October 2015 – May 2017\n"
    "  - Distributed work duties and ensured quality control of cleaning operations.\n"
    "  - Provided hands-on support and training to staff."
)

doc.add_heading("PROJECTS", level=2)
doc.add_paragraph(
    "- FIAM (Ongoing)\n"
    "  - Creating a Python-based program to automate the creation of Flask projects.\n"
    "  - GitHub Repository for FIAM: https://github.com/roldegabell100/FIAM\n"
    "  - Aims to streamline project setup with predefined templates and configurations.\n\n"
    "- TopStock (Ongoing)\n"
    "  - Creating a website that displays the top stocks of the day.\n"
    "  - Features include adding stocks to a personalized watchlist and viewing the most purchased stocks of the day.\n"
    "  - Provides recommendations on stocks to buy and includes links for further information.\n"
    "  - Displays top stories related to stocks for comprehensive market insights.\n"
    "  - GitHub Repository for TopStock: https://github.com/roldegabell100/TopStock\n\n"
    "- Robotics Support\n"
    "  - Designed and built a seven-headed dragon for a theater production.\n"
    "  - Programmed Arduino boards and assembled custom circuit boards.\n"
    "  - GitHub Repository for Robotics Projects: https://github.com/roldegabell100/Arduino"
)

doc.add_heading("AWARDS & RECOGNITION", level=2)
doc.add_paragraph(
    "- Leadership Award, Bear Creek High School\n"
    "- Served as President and Vice President, National Society of Leadership and Success, Red Rocks Community College\n"
    "- Earned two Leadership Certifications from the National Society of Leadership and Success"
)

doc.add_heading("ADDITIONAL PROJECTS", level=2)
doc.add_paragraph("I am currently working on several other projects and can present them if needed.")

doc.add_heading("REFERENCES", level=2)
doc.add_paragraph("Available upon request")

# Save the document to a file
file_path = "/mnt/data/Deconoix_Gloire_Gabell_Resume.docx"
doc.save(file_path)
file_path
